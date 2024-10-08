import model.Banco.BD as conexao
import openpyxl
from docx import Document
import os 
import shutil

#Login - pageLogin
def LoginUsuario(login, senha):
    conexaoBD = conexao.iniciaConexao()
    query = "SELECT * FROM usuarios WHERE login = %s and senha = %s;"
    parametros = (login, senha) 

    cursorBD = conexaoBD.cursor()
    cursorBD.execute(query, parametros)
    listaResultado = cursorBD.fetchone()
    cursorBD.close()
    conexaoBD.close()
    return listaResultado


#Cadastrar Novo Ususario - pageCadastrar
def CadastrarUsuario(login, senha):
    try:
        conexaoBD = conexao.iniciaConexao()
        cursorBD = conexaoBD.cursor()

        cadastrar = "INSERT INTO usuarios (login, senha) VALUES (%s, %s);"
        parametros = (login, senha) 
        cursorBD.execute(cadastrar, parametros)
        conexaoBD.commit()
        cursorBD.close()
        conexaoBD.close()
        return True
    except:
        return False

def DeletarUsuario(id_Usuario):
    conexaoBD = conexao.iniciaConexao()
    query = 'DELETE FROM usuarios WHERE id = %s;'
    parametro = [id_Usuario]
    cursorBD = conexaoBD.cursor()

    cursorBD.execute(query, parametro)
    conexaoBD.commit()
    cursorBD.close()
    conexaoBD.close()



#Ver Usuarios========================================
def ExibirUsuarios():
    listaUsuarios = []
    conexaoBD = conexao.iniciaConexao()
    query = "SELECT * FROM usuarios"

    cursorBD = conexaoBD.cursor()
    cursorBD.execute(query)
    for aluno in cursorBD:
        listaUsuarios.append(aluno)
    cursorBD.close()
    conexaoBD.close()
    return listaUsuarios

#Cadastrar Turma - Tem que mudar umas coisas
def CadastrarTurma(planilha):
    try:
        wb = openpyxl.load_workbook(planilha)
        sheet = wb['3TDSA']

        alunos = []
        for row in sheet.iter_rows(min_row = 2, values_only=True):
            alunos.append(row)

        wb.close()

        for aluno in alunos:
            conexaoBD = conexao.iniciaConexao()
            query = "INSERT INTO bdalunos (instituicao, nome_completo, cpf, rg, orgao_expedidor, municipio, nacionalidade, data_nascimento, curso, data_conclusão, nome_pai, nome_mae, turma) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s);"
            parametro = (aluno[0], aluno[1], aluno[2],aluno[3], aluno[4], aluno[5], aluno[6], aluno[7], (aluno[8]), str(aluno[9]), aluno[10], aluno[11], aluno[12])
            cursorBD = conexaoBD.cursor()
            cursorBD.execute(query, parametro)
            conexaoBD.commit()
            cursorBD.close()
            conexaoBD.close()
        return True
    except:
        return False


#ListarAlunos - 
def ListarAlunos(seletor):
    lista = []
    conexaoBD = conexao.iniciaConexao()
    if seletor == None or seletor == "*":
        query = "SELECT * FROM bdalunos"

        cursorBD = conexaoBD.cursor()
        cursorBD.execute(query)
        for aluno in cursorBD:
            lista.append(aluno)
        cursorBD.close()
        conexaoBD.close()
        return lista
    elif seletor == "3TDSA":
        query = "SELECT * FROM bdalunos WHERE turma = '3TDSA';"

        cursorBD = conexaoBD.cursor()
        cursorBD.execute(query)
        for aluno in cursorBD:
            lista.append(aluno)
        cursorBD.close()
        conexaoBD.close()
        return lista
    
    elif seletor == "3TDSB":
        query = "SELECT * FROM bdalunos WHERE turma = '3TDSB';"

        cursorBD = conexaoBD.cursor()
        cursorBD.execute(query)
        for aluno in cursorBD:
            lista.append(aluno)
        cursorBD.close()
        conexaoBD.close()
        if lista != None:
            return lista
        
    
    elif seletor == "3MKTA":
        query = "SELECT * FROM bdalunos WHERE turma = '3MKTA';"

        cursorBD = conexaoBD.cursor()
        cursorBD.execute(query)
        for aluno in cursorBD:
            lista.append(aluno)
        cursorBD.close()
        conexaoBD.close()
        return lista
    
    elif seletor == "3MKTB":
        query = "SELECT * FROM bdalunos WHERE turma = '3MKTB';"

        cursorBD = conexaoBD.cursor()
        cursorBD.execute(query)
        for aluno in cursorBD:
            lista.append(aluno)
        cursorBD.close()
        conexaoBD.close()
        return lista

#AÇOES
#Exibir Info Alunos -- Editar Alunos
def ExibirAluno(id_Aluno):
    conexaoBD = conexao.iniciaConexao()
    query = "SELECT * FROM bdalunos WHERE id = " + str(id_Aluno) +";"
    cursorBD = conexaoBD.cursor()

    cursorBD.execute(query)
    infoAluno = cursorBD.fetchone()
    cursorBD.close()
    conexaoBD.close()
    return infoAluno 

#Alterar Info Alunos -- Editar Alunos
def EditarAluno(id_Aluno, nome, nomeMae, nomePai, municipioAluno, nacionalidadeAluno, turmaAluno,cpfAluno,rgAluno, cursoAluno):
    conexaoBD = conexao.iniciaConexao()
    query = "UPDATE bdalunos SET nome_completo = %s, nome_mae = %s, nome_pai = %s, municipio = %s, nacionalidade = %s, turma = %s, CPF = %s, rg = %s, curso = %s WHERE id = %s;"
    parametros = (nome, nomeMae, nomePai, municipioAluno, nacionalidadeAluno, turmaAluno, cpfAluno, rgAluno, cursoAluno, id_Aluno)
    cursorBD = conexaoBD.cursor()

    cursorBD.execute(query, parametros)
    conexaoBD.commit()
    cursorBD.close()
    conexaoBD.close()

#Deletar Aluno
def Deletar(id_Aluno):
    conexaoBD = conexao.iniciaConexao()
    query = 'DELETE FROM bdalunos WHERE id = %s;'
    parametro = [id_Aluno]
    cursorBD = conexaoBD.cursor()

    cursorBD.execute(query, parametro)
    conexaoBD.commit()
    cursorBD.close()
    conexaoBD.close()

#CERTIFICADOS ==================================================================================

#Gerar Certificado Individual
def GerarCertificado(documento, id_Aluno):
    conexaoBD = conexao.iniciaConexao()
    query = "SELECT * FROM bdalunos WHERE id = " + str(id_Aluno) +";"
    cursorBD = conexaoBD.cursor()

    cursorBD.execute(query)
    infoAluno = cursorBD.fetchone()
    cursorBD.close()
    conexaoBD.close()
    
    
    # Carregar o documento existente
    doc = Document(documento)
    
    substituicoes = {
        '#instituição#': infoAluno[0],
        '#curso#': infoAluno[10],
        '#nome#':  infoAluno[1],
        '#nomeMae#':  infoAluno[2],
        '#nomePai#':  infoAluno[3],
        '#municipio#':  infoAluno[4],
        '#uf#':  infoAluno[9],
        '#nacionalidade#':  infoAluno[5],
        '#diaNas#':  infoAluno[6],
        '#cpf#':  infoAluno[7],
        '#rg#':  infoAluno[0],
        '#uf#':  infoAluno[9],
        '#diaCon#':  infoAluno[11],
    }


    # Substituir as palavras em parágrafos
    for para in doc.paragraphs:
        for palavra_antiga, palavra_nova in substituicoes.items():
            if palavra_antiga in para.text: 
                para.text = para.text.replace(palavra_antiga, palavra_nova)
                


    # Salvar o documento modificado]
    pastaCertificados = os.path.abspath("static/Certificados/" + str(infoAluno[12]) + "-" + infoAluno[1] +".docx")
    doc.save(pastaCertificados)


def GerarTodosCertificados(seletor, documento):
    listaAlunos = ListarAlunos(seletor)
    
    for aluno in listaAlunos:
        doc = Document(documento)
        substituicoes = {
            '#instituição#': listaAlunos[0],
            '#curso#': listaAlunos[10],
            '#nome#':  listaAlunos[1],
            '#nomeMae#':  listaAlunos[2],
            '#nomePai#':  listaAlunos[3],
            '#municipio#':  listaAlunos[4],
            '#uf#':  listaAlunos[9],
            '#nacionalidade#':  listaAlunos[5],
            '#diaNas#':  listaAlunos[6],
            '#cpf#':  listaAlunos[7],
            '#rg#':  listaAlunos[0],
            '#uf#':  listaAlunos[9],
            '#diaCon#':  listaAlunos[11],
        }


        # Substituir as palavras em parágrafos
        for para in doc.paragraphs:
            for palavra_antiga, palavra_nova in substituicoes.items():
                if palavra_antiga in para.text: 
                    para.text = para.text.replace(palavra_antiga, palavra_nova)
    

        # Salvar o documento modificado]
        pastaCertificados = os.path.abspath("static/Certificados/" + aluno[1] + ".docx")
        doc.save(pastaCertificados)
        
    shutil.make_archive(os.path.abspath("static/Certificados/Certificados"), 'zip', os.path.abspath("static/Certificados/"))
    return "Certificados.zip"


def certificados(certificadosCriados):
    lista = []
    conexaoBD = conexao.iniciaConexao()
    cursorBD = conexaoBD.cursor()
    
    for idCertificado in certificadosCriados:
        query = "SELECT * FROM bdalunos WHERE id = " + str(idCertificado)
        cursorBD.execute(query)
        lista.append(cursorBD.fetchone())

    cursorBD.close()
    conexaoBD.close()
    return lista

#Açoes pageCertificados---